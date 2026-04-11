"""
CCAP Analyzer - Service d'extraction de texte
Lexigency - 2026

Extrait le texte des fichiers DOCX/PDF et charge le Code de la Commande Publique.
Les PDF peuvent être convertis en DOCX pour permettre l'annotation Word native.
"""

import logging
from pathlib import Path
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from config import (
    CODE_COMMANDE_PUBLIQUE_FILE,
    MAX_CHARS_CCAG,
    MAX_CHARS_CCTP,
    MAX_CHARS_CODE_CCP,
)

logger = logging.getLogger(__name__)

# =============================================================================
# EXTRACTION DE TEXTE DOCX
# =============================================================================

def extract_text_from_docx(docx_path, max_chars=None):
    """
    Extrait le texte d'un fichier DOCX.

    Args:
        docx_path: Chemin vers le fichier DOCX (str ou Path)
        max_chars: Nombre maximum de caractères à extraire (None = pas de limite)

    Returns:
        str: Texte extrait du document
    """
    docx_path = Path(docx_path)

    if not docx_path.exists():
        raise FileNotFoundError(f"Fichier introuvable: {docx_path}")

    if not docx_path.suffix.lower() == ".docx":
        raise ValueError(f"Extension invalide: {docx_path.suffix}. Attendu: .docx")

    try:
        doc = Document(docx_path)
    except PackageNotFoundError:
        raise ValueError(f"Fichier DOCX invalide ou corrompu: {docx_path}")
    except Exception as e:
        raise ValueError(f"Erreur lors de l'ouverture du fichier: {e}")

    text_parts = []

    # Extraire le texte des paragraphes
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            text_parts.append(text)

    # Extraire le texte des tableaux
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                text_parts.append(" | ".join(row_text))

    full_text = "\n".join(text_parts)

    if max_chars and len(full_text) > max_chars:
        logger.warning(f"Texte tronqué de {len(full_text)} à {max_chars} caractères")
        full_text = full_text[:max_chars]
        last_period = full_text.rfind(".")
        last_newline = full_text.rfind("\n")
        cut_point = max(last_period, last_newline)
        if cut_point > max_chars * 0.8:
            full_text = full_text[: cut_point + 1]

    logger.info(f"Texte extrait: {len(full_text)} caractères depuis {docx_path.name}")
    return full_text


def extract_ccag(docx_path):
    """Extrait le texte d'un fichier CCAG."""
    return extract_text_from_docx(docx_path, max_chars=MAX_CHARS_CCAG)


def extract_cctp(docx_path):
    """Extrait le texte d'un fichier CCTP."""
    return extract_text_from_docx(docx_path, max_chars=MAX_CHARS_CCTP)


# =============================================================================
# EXTRACTION DE TEXTE PDF
# =============================================================================

def extract_text_from_pdf(pdf_path, max_chars=None):
    """
    Extrait le texte d'un fichier PDF via PyMuPDF (fitz).

    Args:
        pdf_path: Chemin vers le fichier PDF (str ou Path)
        max_chars: Nombre maximum de caractères à extraire (None = pas de limite)

    Returns:
        str: Texte brut extrait du document
    """
    pdf_path = Path(pdf_path)

    if not pdf_path.exists():
        raise FileNotFoundError(f"Fichier introuvable: {pdf_path}")

    if not pdf_path.suffix.lower() == ".pdf":
        raise ValueError(f"Extension invalide: {pdf_path.suffix}. Attendu: .pdf")

    try:
        import fitz  # PyMuPDF
    except ImportError as e:
        raise ImportError(
            "PyMuPDF est requis pour l'extraction PDF. Installez avec: pip install PyMuPDF"
        ) from e

    try:
        doc = fitz.open(pdf_path)
    except Exception as e:
        raise ValueError(f"Fichier PDF invalide ou corrompu: {pdf_path} ({e})")

    text_parts = []
    try:
        for page in doc:
            page_text = page.get_text("text")
            if page_text:
                text_parts.append(page_text.strip())
    finally:
        doc.close()

    full_text = "\n".join(text_parts)

    if max_chars and len(full_text) > max_chars:
        logger.warning(f"Texte PDF tronqué de {len(full_text)} à {max_chars} caractères")
        full_text = full_text[:max_chars]
        last_period = full_text.rfind(".")
        last_newline = full_text.rfind("\n")
        cut_point = max(last_period, last_newline)
        if cut_point > max_chars * 0.8:
            full_text = full_text[: cut_point + 1]

    logger.info(f"Texte PDF extrait: {len(full_text)} caractères depuis {pdf_path.name}")
    return full_text


# =============================================================================
# CONVERSION PDF → DOCX
# =============================================================================

def convert_pdf_to_docx(pdf_path):
    """
    Convertit un fichier PDF en fichier DOCX via pdf2docx.

    Le .docx résultant est sauvegardé à côté du PDF original, avec le même
    nom de base et l'extension .docx. Si un fichier .docx du même nom existe
    déjà, il est écrasé.

    Args:
        pdf_path: Chemin vers le fichier PDF (str ou Path)

    Returns:
        Path: Chemin du fichier .docx converti
    """
    pdf_path = Path(pdf_path)

    if not pdf_path.exists():
        raise FileNotFoundError(f"Fichier PDF introuvable: {pdf_path}")

    if not pdf_path.suffix.lower() == ".pdf":
        raise ValueError(f"Extension invalide: {pdf_path.suffix}. Attendu: .pdf")

    try:
        from pdf2docx import Converter
    except ImportError as e:
        raise ImportError(
            "pdf2docx est requis pour la conversion PDF → DOCX. "
            "Installez avec: pip install pdf2docx"
        ) from e

    docx_path = pdf_path.with_suffix(".docx")

    logger.info(f"Conversion PDF → DOCX: {pdf_path.name} → {docx_path.name}")

    try:
        cv = Converter(str(pdf_path))
        try:
            cv.convert(str(docx_path))
        finally:
            cv.close()
    except Exception as e:
        raise ValueError(f"Erreur lors de la conversion PDF → DOCX: {e}")

    if not docx_path.exists():
        raise RuntimeError(f"La conversion PDF → DOCX a échoué (fichier non créé): {docx_path}")

    logger.info(f"Conversion réussie: {docx_path}")
    return docx_path


# =============================================================================
# DISPATCHER : EXTRACTION SELON L'EXTENSION
# =============================================================================

def extract_text(filepath, max_chars=None):
    """
    Extrait le texte d'un fichier DOCX ou PDF selon son extension.

    Args:
        filepath: Chemin vers le fichier (str ou Path) — .docx ou .pdf
        max_chars: Nombre maximum de caractères à extraire

    Returns:
        str: Texte extrait du document
    """
    filepath = Path(filepath)
    suffix = filepath.suffix.lower()

    if suffix == ".docx":
        return extract_text_from_docx(filepath, max_chars=max_chars)
    elif suffix == ".pdf":
        return extract_text_from_pdf(filepath, max_chars=max_chars)
    else:
        raise ValueError(f"Extension non supportée: {suffix}. Attendu: .docx ou .pdf")


# =============================================================================
# CHARGEMENT DU CODE DE LA COMMANDE PUBLIQUE
# =============================================================================

_code_ccp_cache = None

def load_code_commande_publique():
    """Charge le Code de la Commande Publique depuis le fichier fixe (avec cache)."""
    global _code_ccp_cache

    if _code_ccp_cache is not None:
        return _code_ccp_cache

    if not CODE_COMMANDE_PUBLIQUE_FILE.exists():
        raise FileNotFoundError(
            f"Fichier Code de la Commande Publique introuvable: {CODE_COMMANDE_PUBLIQUE_FILE}"
        )

    try:
        with open(CODE_COMMANDE_PUBLIQUE_FILE, "r", encoding="utf-8") as f:
            full_text = f.read()
    except UnicodeDecodeError:
        with open(CODE_COMMANDE_PUBLIQUE_FILE, "r", encoding="latin-1") as f:
            full_text = f.read()

    if len(full_text) > MAX_CHARS_CODE_CCP:
        logger.warning(
            f"Code CCP tronqué de {len(full_text)} à {MAX_CHARS_CODE_CCP} caractères"
        )
        full_text = full_text[:MAX_CHARS_CODE_CCP]
        last_article = full_text.rfind("\nArticle")
        if last_article > MAX_CHARS_CODE_CCP * 0.8:
            full_text = full_text[:last_article]

    _code_ccp_cache = full_text
    logger.info(f"Code de la Commande Publique chargé: {len(full_text)} caractères")
    return full_text


def clear_code_ccp_cache():
    """Vide le cache du Code de la Commande Publique."""
    global _code_ccp_cache
    _code_ccp_cache = None


# =============================================================================
# FONCTIONS UTILITAIRES
# =============================================================================

def get_document_info(docx_path):
    """Retourne des informations sur un document DOCX."""
    docx_path = Path(docx_path)

    if not docx_path.exists():
        return {"error": "Fichier introuvable"}

    try:
        doc = Document(docx_path)
        num_paragraphs = len(doc.paragraphs)
        num_tables = len(doc.tables)
        text = extract_text_from_docx(docx_path)
        num_chars = len(text)
        num_words = len(text.split())

        return {
            "filename": docx_path.name,
            "size_bytes": docx_path.stat().st_size,
            "paragraphs": num_paragraphs,
            "tables": num_tables,
            "characters": num_chars,
            "words": num_words,
        }
    except Exception as e:
        return {"error": str(e)}
