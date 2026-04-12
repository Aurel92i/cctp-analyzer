"""
CCAP Analyzer - Service d'annotation Word avec COMMENTAIRES NATIFS
Lexigency - 2026

Ce module insère de VRAIS commentaires Word (panneau Révisions)
dans le CCAP à partir des remarques générées par l'analyse.

Les commentaires apparaissent dans le panneau latéral de Word,
comme des commentaires de révision classiques.
"""

import logging
import re
import zipfile
import shutil
import os
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from lxml import etree

logger = logging.getLogger(__name__)

# =============================================================================
# CONSTANTES
# =============================================================================

COMMENT_AUTHOR = "LEXIGENCY"
COMMENT_INITIALS = "LX"

GRAVITY_HEADERS = {
    "haute": "🔴 GRAVITÉ HAUTE",
    "moyenne": "🟠 GRAVITÉ MOYENNE",
    "basse": "🟢 GRAVITÉ BASSE",
}

# Namespaces XML pour Word
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"

WORD_NSMAP = {
    "w": W_NS,
    "r": R_NS,
}


# =============================================================================
# FONCTION PRINCIPALE
# =============================================================================

def annotate_document(ccap_path: str, remarques: list, output_path: str) -> dict:
    """
    Annote un CCAP avec de VRAIS commentaires Word natifs.

    Args:
        ccap_path: Chemin vers le CCAP original (document à annoter)
        remarques: Liste des remarques de l'analyse
        output_path: Chemin de sortie pour le CCAP annoté

    Returns:
        dict: Statistiques d'annotation
    """
    try:
        logger.info(f"Annotation du CCAP: {ccap_path}")
        logger.info(f"Nombre de remarques: {len(remarques)}")

        if not remarques:
            shutil.copy(ccap_path, output_path)
            return {
                "success": True,
                "output_path": str(output_path),
                "comments_added": 0,
                "comments_not_found": 0,
                "total_remarques": 0,
            }

        # Créer l'annotateur
        annotator = WordCommentAnnotator(ccap_path)

        # Ajouter chaque remarque comme commentaire
        comments_added = 0
        comments_not_found = 0

        for i, remarque in enumerate(remarques):
            success = annotator.add_comment(remarque, comment_id=i)
            if success:
                comments_added += 1
            else:
                comments_not_found += 1

        # Sauvegarder le document
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        annotator.save(output_path)

        logger.info(f"CCAP annoté sauvegardé: {output_path}")
        logger.info(f"Commentaires ajoutés: {comments_added}/{len(remarques)}")

        return {
            "success": True,
            "output_path": str(output_path),
            "comments_added": comments_added,
            "comments_not_found": comments_not_found,
            "total_remarques": len(remarques),
        }

    except Exception as e:
        logger.error(f"Erreur annotation: {e}")
        import traceback

        traceback.print_exc()
        return {
            "success": False,
            "error": str(e),
            "comments_added": 0,
            "comments_not_found": 0,
        }


# =============================================================================
# CLASSE D'ANNOTATION AVEC COMMENTAIRES NATIFS
# =============================================================================

class WordCommentAnnotator:
    """
    Ajoute des commentaires Word natifs à un document CCAP.
    Manipule directement le XML OOXML pour insérer des commentaires
    qui apparaissent dans le panneau Révisions de Word.
    """

    def __init__(self, docx_path: str):
        self.docx_path = Path(docx_path)
        self.doc = Document(docx_path)
        self.comments = []
        self.next_comment_id = 0

    def _format_comment_text(self, remarque: dict) -> str:
        """Formate le texte du commentaire à partir d'une remarque."""
        gravite = remarque.get("gravite", "moyenne")
        header = GRAVITY_HEADERS.get(gravite, "🟠 GRAVITÉ MOYENNE")

        parts = []

        # Préfixe pour les remarques CCTP
        if remarque.get("document_source") == "CCTP":
            parts.append("📄 REMARQUE SUR LE CCTP")
            section = remarque.get("section_source", "")
            if section:
                parts.append(f"Section : {section}")
            parts.append("")

        parts.extend([
            header,
            "━" * 25,
            "",
            "📋 CONSTAT:",
            remarque.get("constat", "N/A"),
            "",
            "⚠️ PROBLÈME:",
            remarque.get("probleme", "N/A"),
            "",
            "📖 RÉFÉRENCES:",
            remarque.get("references_juridiques", "N/A"),
            "",
            "✅ RECOMMANDATION:",
            remarque.get("recommandation", "N/A"),
        ])

        return "\n".join(parts)

    def _normalize_text(self, text: str) -> str:
        """Normalise un texte pour la comparaison."""
        if not text:
            return ""
        return " ".join(text.lower().split())

    def _iter_all_paragraphs(self):
        """Itère sur tous les paragraphes du document (corps + tableaux)."""
        for para in self.doc.paragraphs:
            yield para
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        yield para

    def _find_paragraph_with_text(self, search_text: str):
        """
        Trouve le paragraphe contenant le texte recherché.
        Utilise un matching flou (SequenceMatcher) avec threshold 0.65.
        """
        from difflib import SequenceMatcher

        if not search_text or len(search_text) < 10:
            return None, False

        search_normalized = self._normalize_text(search_text)

        # Stratégie 1 : recherche exacte
        for para in self._iter_all_paragraphs():
            para_normalized = self._normalize_text(para.text)
            if search_normalized in para_normalized:
                return para, True

        # Stratégie 2 : recherche par les 50 premiers caractères
        search_start = search_normalized[:50]
        for para in self._iter_all_paragraphs():
            para_normalized = self._normalize_text(para.text)
            if len(para_normalized) > 20 and search_start in para_normalized:
                return para, True

        # Stratégie 3 : matching flou
        best_match = None
        best_score = 0
        threshold = 0.65

        for para in self._iter_all_paragraphs():
            para_normalized = self._normalize_text(para.text)
            if len(para_normalized) < 20:
                continue

            if len(para_normalized) > len(search_normalized) * 2:
                window = len(search_normalized)
                for start in range(0, len(para_normalized) - window + 1, 20):
                    chunk = para_normalized[start : start + window]
                    score = SequenceMatcher(
                        None, search_normalized[:100], chunk[:100]
                    ).ratio()
                    if score > best_score:
                        best_score = score
                        best_match = para
            else:
                score = SequenceMatcher(
                    None, search_normalized[:100], para_normalized[:100]
                ).ratio()
                if score > best_score:
                    best_score = score
                    best_match = para

        if best_match and best_score >= threshold:
            return best_match, True

        return None, False

    def _create_comment_element(self, comment_id: int, text: str) -> etree._Element:
        """Crée l'élément XML du commentaire."""
        comment = OxmlElement("w:comment")
        comment.set(qn("w:id"), str(comment_id))
        comment.set(qn("w:author"), COMMENT_AUTHOR)
        comment.set(qn("w:initials"), COMMENT_INITIALS)
        comment.set(qn("w:date"), datetime.now().strftime("%Y-%m-%dT%H:%M:%S"))

        for line in text.split("\n"):
            p = OxmlElement("w:p")
            r = OxmlElement("w:r")
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = line
            r.append(t)
            p.append(r)
            comment.append(p)

        return comment

    def _add_comment_to_paragraph(self, para, comment_id: int, comment_text: str):
        """Ajoute les marqueurs de commentaire à un paragraphe."""
        p = para._element

        range_start = OxmlElement("w:commentRangeStart")
        range_start.set(qn("w:id"), str(comment_id))

        range_end = OxmlElement("w:commentRangeEnd")
        range_end.set(qn("w:id"), str(comment_id))

        comment_ref_run = OxmlElement("w:r")
        comment_ref = OxmlElement("w:commentReference")
        comment_ref.set(qn("w:id"), str(comment_id))
        comment_ref_run.append(comment_ref)

        if len(p) > 0:
            p.insert(0, range_start)
        else:
            p.append(range_start)

        p.append(range_end)
        p.append(comment_ref_run)

        comment_element = self._create_comment_element(comment_id, comment_text)
        self.comments.append(comment_element)

    def _get_last_paragraph(self):
        """Retourne le dernier paragraphe non-vide du document."""
        paragraphs = list(self._iter_all_paragraphs())
        for para in reversed(paragraphs):
            if para.text.strip():
                return para
        return paragraphs[-1] if paragraphs else None

    def add_comment(self, remarque: dict, comment_id: int = None) -> bool:
        """Ajoute un commentaire pour une remarque."""
        extrait = remarque.get("extrait_texte", "")

        # Remarques CCTP : placer sur le dernier paragraphe du document CCAP
        if remarque.get("document_source") == "CCTP":
            para = self._get_last_paragraph()
            if not para:
                logger.warning("Aucun paragraphe dans le document pour le commentaire CCTP")
                return False
            found = True
        else:
            if not extrait:
                logger.warning("Remarque sans extrait de texte, ignorée")
                return False

            para, found = self._find_paragraph_with_text(extrait)

        if not found:
            logger.warning(f"Texte non trouvé dans le CCAP: {extrait[:80]}...")
            return False

        if comment_id is None:
            comment_id = self.next_comment_id
        self.next_comment_id = max(self.next_comment_id, comment_id + 1)

        comment_text = self._format_comment_text(remarque)
        self._add_comment_to_paragraph(para, comment_id, comment_text)

        logger.debug(f"Commentaire {comment_id} ajouté pour: {extrait[:50]}...")
        return True

    def _inject_comments_part(self, docx_path: str):
        """Injecte la partie comments.xml dans le document DOCX."""
        if not self.comments:
            return

        comments_root = etree.Element(qn("w:comments"), nsmap={"w": W_NS})
        for comment in self.comments:
            comments_root.append(comment)

        comments_xml = etree.tostring(
            comments_root,
            xml_declaration=True,
            encoding="UTF-8",
            standalone=True,
        )

        with zipfile.ZipFile(docx_path, "a") as zf:
            zf.writestr("word/comments.xml", comments_xml)

        self._update_content_types(docx_path)
        self._update_relationships(docx_path)

    def _update_content_types(self, docx_path: str):
        """Met à jour [Content_Types].xml pour inclure comments.xml."""
        try:
            with zipfile.ZipFile(docx_path, "r") as zf:
                content_types = zf.read("[Content_Types].xml").decode("utf-8")

            if "word/comments.xml" not in content_types:
                override = '<Override PartName="/word/comments.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/>'
                content_types = content_types.replace(
                    "</Types>", f"{override}\n</Types>"
                )
                self._update_zip_file(
                    docx_path, "[Content_Types].xml", content_types.encode("utf-8")
                )
        except Exception as e:
            logger.warning(f"Erreur mise à jour Content_Types: {e}")

    def _update_relationships(self, docx_path: str):
        """Met à jour word/_rels/document.xml.rels pour inclure comments.xml."""
        try:
            rels_path = "word/_rels/document.xml.rels"

            with zipfile.ZipFile(docx_path, "r") as zf:
                if rels_path in zf.namelist():
                    rels_content = zf.read(rels_path).decode("utf-8")
                else:
                    rels_content = '<?xml version="1.0" encoding="UTF-8"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"></Relationships>'

            if "comments.xml" not in rels_content:
                existing_ids = re.findall(r'Id="rId(\d+)"', rels_content)
                next_id = max([int(x) for x in existing_ids], default=0) + 1

                rel = f'<Relationship Id="rId{next_id}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" Target="comments.xml"/>'
                rels_content = rels_content.replace(
                    "</Relationships>", f"{rel}\n</Relationships>"
                )
                self._update_zip_file(
                    docx_path, rels_path, rels_content.encode("utf-8")
                )
        except Exception as e:
            logger.warning(f"Erreur mise à jour Relations: {e}")

    def _update_zip_file(self, zip_path: str, file_name: str, content: bytes):
        """Met à jour un fichier dans un ZIP existant."""
        temp_path = zip_path + ".tmp"

        with zipfile.ZipFile(zip_path, "r") as zin:
            with zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as zout:
                for item in zin.namelist():
                    if item == file_name:
                        zout.writestr(item, content)
                    else:
                        zout.writestr(item, zin.read(item))

                if file_name not in zin.namelist():
                    zout.writestr(file_name, content)

        os.replace(temp_path, zip_path)

    def save(self, output_path: str):
        """Sauvegarde le document CCAP annoté."""
        self.doc.save(output_path)

        if self.comments:
            self._inject_comments_part(output_path)

        logger.info(
            f"CCAP sauvegardé avec {len(self.comments)} commentaires: {output_path}"
        )
